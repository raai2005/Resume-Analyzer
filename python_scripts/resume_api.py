"""
Resume Analysis API Endpoint - Complete feedback report generation
"""

import json
import sys
import os
from typing import Dict, Any
from document_parser import parse_document
from feedback_report import FeedbackReportGenerator


def analyze_resume_complete(file_path: str, 
                          job_description: str = None,
                          job_title: str = None,
                          company: str = None,
                          required_skills: list = None,
                          preferred_skills: list = None) -> Dict[str, Any]:
    """
    Complete resume analysis with structured feedback report
    
    Args:
        file_path: Path to the resume file
        job_description: Optional job description for targeted analysis
        job_title: Optional job title for context
        company: Optional company name
        required_skills: Optional list of required skills
        preferred_skills: Optional list of preferred skills
        
    Returns:
        Structured feedback report with 200 OK response format
    """
    
    try:
        # Step 1: Parse the document with all analysis steps
        print(f"🔍 Analyzing resume: {os.path.basename(file_path)}")
        
        parse_result = parse_document(
            file_path=file_path,
            job_description=job_description,
            job_title=job_title,
            company=company,
            required_skills=required_skills,
            preferred_skills=preferred_skills
        )
        
        if not parse_result.get("success"):
            return {
                "status": "error",
                "status_code": 400,
                "error": parse_result.get("error", "Document parsing failed"),
                "details": parse_result
            }
        
        # Step 2: Generate structured feedback report
        print("📊 Generating structured feedback report...")
        
        generator = FeedbackReportGenerator()
        feedback_report = generator.generate_feedback_report(
            parse_result=parse_result,
            job_description=job_description,
            required_skills=required_skills
        )
        
        # Step 3: Format as API response with 200 OK structure
        api_response = {
            "status": "success",
            "status_code": 200,
            "message": "Resume analysis completed successfully",
            "data": feedback_report
        }
        
        print("✅ Analysis completed successfully!")
        return api_response
        
    except Exception as e:
        print(f"❌ Analysis failed: {str(e)}")
        return {
            "status": "error", 
            "status_code": 500,
            "error": f"Internal server error: {str(e)}",
            "details": {}
        }


def test_complete_analysis():
    """Test the complete analysis with a sample resume"""
    
    print("Testing Complete Resume Analysis API")
    print("=" * 50)
    
    # Create a test resume file
    from docx import Document
    
    doc = Document()
    doc.add_heading('ALEX CHEN', 0)
    doc.add_paragraph('Senior Software Engineer')
    doc.add_paragraph('Email: alex.chen@email.com | Phone: (555) 123-4567')
    doc.add_paragraph('LinkedIn: linkedin.com/in/alexchen | GitHub: github.com/alexchen')
    doc.add_paragraph('')
    
    doc.add_heading('EXPERIENCE', level=1)
    doc.add_paragraph('Senior Software Engineer - TechCorp (2022-2024)')
    doc.add_paragraph('• Led development of microservices platform serving 100,000+ users')
    doc.add_paragraph('• Implemented React/TypeScript frontend reducing load time by 40%')
    doc.add_paragraph('• Mentored team of 5 junior developers')
    doc.add_paragraph('')
    
    doc.add_paragraph('Software Engineer - StartupXYZ (2020-2022)')
    doc.add_paragraph('• Built RESTful APIs with Node.js handling 5,000+ requests/minute')
    doc.add_paragraph('• Developed automated testing suite achieving 85% code coverage')
    doc.add_paragraph('')
    
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    doc.add_paragraph('University of California, Berkeley (2016-2020)')
    doc.add_paragraph('')
    
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph('Programming: JavaScript, TypeScript, Python, Java')
    doc.add_paragraph('Frontend: React, Vue.js, HTML5, CSS3')
    doc.add_paragraph('Backend: Node.js, Express.js, Django')
    doc.add_paragraph('Databases: PostgreSQL, MongoDB, Redis')
    doc.add_paragraph('Tools: Docker, AWS, Git, Jenkins')
    
    test_file = "test_api_resume.docx"
    doc.save(test_file)
    
    try:
        # Test the complete analysis
        result = analyze_resume_complete(
            file_path=test_file,
            job_description="""
            We are seeking a Senior Software Engineer to join our team.
            
            Requirements:
            - 5+ years of software development experience
            - Strong JavaScript/TypeScript and React skills
            - Experience with Node.js and API development
            - Database experience (PostgreSQL preferred)
            - Cloud platform experience (AWS)
            - Leadership and mentoring experience
            
            Nice to have:
            - Python experience
            - Docker/containerization
            - CI/CD experience
            """,
            job_title="Senior Software Engineer",
            company="TechInnovate",
            required_skills=["JavaScript", "TypeScript", "React", "Node.js", "PostgreSQL", "AWS"],
            preferred_skills=["Python", "Docker", "Leadership", "Mentoring"]
        )
        
        print(f"📋 API Response Status: {result['status']} ({result['status_code']})")
        
        if result['status'] == 'success':
            data = result['data']
            
            # Display key information from the structured response
            print(f"\n👤 Contact Information:")
            contact = data['contact_info']
            print(f"  Name: {contact['name']}")
            print(f"  Email: {contact['email']} (Valid: {contact['email_valid']})")
            print(f"  Phone: {contact['phone']} (Provided: {contact['phone_provided']})")
            print(f"  Links: {len(contact['links'])} professional links")
            print(f"  Completeness: {contact['completeness_score']}/100")
            
            print(f"\n🎓 Education ({len(data['education'])} entries):")
            for edu in data['education']:
                print(f"  • {edu['degree']} - {edu['institution']} ({edu['year']})")
            
            print(f"\n💼 Experience ({len(data['experience'])} positions):")
            for exp in data['experience']:
                print(f"  • {exp['role']} at {exp['company']} ({exp['duration']}) - {exp['type']}")
            
            print(f"\n🔧 Skills Analysis:")
            skills = data['skills']
            print(f"  Total Skills: {skills['total_count']}")
            print(f"  Skills Matched: {len(skills['matched'])}/{len(skills['matched']) + len(skills['missing'])}")
            print(f"  Match Percentage: {skills['match_percentage']:.1f}%")
            print(f"  Missing Skills: {', '.join(skills['missing'])}")
            print(f"  Bonus Skills: {', '.join(skills['bonus'][:3])}")
            
            print(f"\n📈 Quality Scores:")
            quality = data['quality_scores']
            if quality['available']:
                print(f"  Overall Score: {quality['overall_score']}/100 ({quality['quality_level']})")
                print(f"  Content Fit: {quality['percentiles']['content_fit']:.1f}%")
                print(f"  Clarity: {quality['percentiles']['clarity']:.1f}%")
                print(f"  Structure: {quality['percentiles']['structure']:.1f}%")
                print(f"  ATS Friendliness: {quality['percentiles']['ats_friendliness']:.1f}%")
            
            print(f"\n🤖 ATS Compatibility:")
            ats = data['ats_compatibility']
            if ats['available']:
                print(f"  ATS Score: {ats['score']}/100 ({ats['compatibility_level']})")
                print(f"  Priority Issues: {len(ats['priority_issues'])}")
            
            print(f"\n🎯 Role Inference: {data['role_inference']}")
            
            print(f"\n💪 Key Strengths:")
            for strength in data['strengths']:
                print(f"  • {strength}")
            
            print(f"\n🔧 Improvement Areas:")
            for improvement in data['improvement_areas']:
                print(f"  • {improvement}")
            
            print(f"\n💡 Top Recommendations:")
            for i, rec in enumerate(data['recommendations']['top_3'], 1):
                print(f"  {i}. {rec}")
            
            print(f"\n📊 Document Metrics:")
            metrics = data['document_metrics']
            content = metrics['content_metrics']
            print(f"  Words: {content['total_words']}, Characters: {content['total_characters']}")
            print(f"  Estimated Pages: {content['estimated_pages']}")
            print(f"  Sections: {content['sections_count']}, Bullet Points: {content['bullet_points']}")
            
            print(f"\n🎯 Match Analysis:")
            match = data['match_analysis']
            if match['available']:
                print(f"  Overall Match: {match['overall_match']}/100")
                print(f"  Experience Level: {match['experience_level_match']}")
                print(f"  Skills Coverage: {match['skills_coverage']:.1f}%")
            
            # Show complete JSON structure (truncated for readability)
            print(f"\n📋 Complete API Response Keys:")
            print(f"  Status: {result['status']}")
            print(f"  Data Keys: {list(data.keys())}")
            
            # Save sample response
            with open("sample_api_response.json", "w") as f:
                json.dump(result, f, indent=2, ensure_ascii=False)
            print(f"  💾 Full response saved to: sample_api_response.json")
        
        else:
            print(f"❌ Analysis failed: {result['error']}")
            
    except Exception as e:
        print(f"❌ Test failed: {str(e)}")
        import traceback
        traceback.print_exc()
    
    finally:
        # Clean up test file
        if os.path.exists(test_file):
            os.remove(test_file)
            print(f"🗑️ Cleaned up test file: {test_file}")
    
    print("\n" + "=" * 50)
    print("Complete Analysis API test finished!")


def main():
    """Main function for command line usage"""
    
    if len(sys.argv) < 2:
        print("Resume Analysis API - Usage Examples:")
        print("1. Basic analysis:")
        print("   python resume_api.py <file_path>")
        print("2. With job context:")
        print("   python resume_api.py <file_path> <job_description> <job_title> <company>")
        print("3. With skills:")
        print('   python resume_api.py <file_path> "job desc" "title" "company" \'["skill1","skill2"]\' \'["pref1","pref2"]\'')
        print("4. Run test:")
        print("   python resume_api.py test")
        sys.exit(1)
    
    if sys.argv[1] == "test":
        test_complete_analysis()
        return
    
    file_path = sys.argv[1]
    
    # Parse optional parameters
    job_description = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2] != 'None' else None
    job_title = sys.argv[3] if len(sys.argv) > 3 and sys.argv[3] != 'None' else None
    company = sys.argv[4] if len(sys.argv) > 4 and sys.argv[4] != 'None' else None
    
    # Parse skills arrays
    required_skills = None
    preferred_skills = None
    
    if len(sys.argv) > 5 and sys.argv[5] != 'None':
        try:
            required_skills = json.loads(sys.argv[5])
        except json.JSONDecodeError:
            pass
    
    if len(sys.argv) > 6 and sys.argv[6] != 'None':
        try:
            preferred_skills = json.loads(sys.argv[6])
        except json.JSONDecodeError:
            pass
    
    # Run analysis
    result = analyze_resume_complete(
        file_path=file_path,
        job_description=job_description,
        job_title=job_title,
        company=company,
        required_skills=required_skills,
        preferred_skills=preferred_skills
    )
    
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
