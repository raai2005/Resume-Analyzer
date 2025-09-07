"""
Debug script to identify the slice error
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from document_parser import parse_document
from docx import Document

def create_simple_test_resume():
    """Create a simple test resume"""
    doc = Document()
    doc.add_heading('TEST USER', 0)
    doc.add_paragraph('Software Engineer')
    doc.add_paragraph('test@email.com | (555) 123-4567')
    
    doc.add_heading('EXPERIENCE', level=1)
    doc.add_paragraph('Software Engineer - Company (2022-2024)')
    doc.add_paragraph('• Developed applications')
    
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph('Python, JavaScript, React')
    
    test_file = "debug_test.docx"
    doc.save(test_file)
    return test_file

def debug_parse():
    """Debug the parsing to find the slice error"""
    
    test_file = None
    try:
        test_file = create_simple_test_resume()
        print(f"Created test file: {test_file}")
        
        print("Starting document parsing...")
        result = parse_document(
            file_path=test_file,
            job_description="Software engineer position requiring Python and JavaScript",
            required_skills=["Python", "JavaScript"]
        )
        
        print(f"Parse result success: {result.get('success')}")
        
        if result.get("success"):
            print("Parsing successful, now testing feedback report...")
            
            from feedback_report import FeedbackReportGenerator
            generator = FeedbackReportGenerator()
            
            print("Calling generate_feedback_report...")
            feedback = generator.generate_feedback_report(
                parse_result=result,
                required_skills=["Python", "JavaScript"]
            )
            
            print("Feedback report generated successfully!")
            print(f"Status: {feedback.get('status')}")
            
        else:
            print(f"Parse failed: {result.get('error')}")
    
    except Exception as e:
        print(f"Error occurred: {str(e)}")
        import traceback
        traceback.print_exc()
    
    finally:
        if test_file and os.path.exists(test_file):
            os.remove(test_file)
            print(f"Cleaned up: {test_file}")

if __name__ == "__main__":
    debug_parse()
