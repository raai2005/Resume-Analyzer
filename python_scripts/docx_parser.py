"""
DOCX text extraction for resume parsing
Uses python-docx library for Word document parsing
"""

import json
import sys
import os
from typing import Dict, Any
from docx import Document


def extract_docx_text(file_path: str) -> Dict[str, Any]:
    """
    Extract text from DOCX using python-docx
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Dictionary containing extraction results
    """
    result = {
        "success": False,
        "text": "",
        "metadata": {},
        "parsing_method": "python-docx",
        "char_count": 0,
        "error": None
    }
    
    try:
        if not os.path.exists(file_path):
            result["error"] = f"File not found: {file_path}"
            return result
            
        # Load the document
        doc = Document(file_path)
        
        # Extract text from paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
        
        # Combine all text
        all_text = "\n".join(text_parts)
        if table_text:
            all_text += "\n\nTables:\n" + "\n".join(table_text)
        
        # Extract metadata
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "category": core_props.category or "",
            "comments": core_props.comments or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision if core_props.revision else 0,
            "version": core_props.version or "",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size": os.path.getsize(file_path)
        }
        
        result.update({
            "success": True,
            "text": all_text,
            "metadata": metadata,
            "char_count": len(all_text)
        })
        
    except Exception as e:
        result["error"] = f"DOCX parsing error: {str(e)}"
        
    return result


def main():
    """Main function for command line usage"""
    if len(sys.argv) != 2:
        print(json.dumps({
            "success": False,
            "error": "Usage: python docx_parser.py <docx_file_path>"
        }))
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = extract_docx_text(file_path)
    
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
