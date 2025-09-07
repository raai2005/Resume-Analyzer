"""
Resume Analyzer - Complete Feedback Report API

Returns structured JSON with:
- 200 OK: Successful analysis with complete feedback report
- 400 Bad Request: Invalid input or file processing error  
- 500 Internal Server Error: System error during analysis

Response Structure:
{
  "status": "success|error",
  "status_code": 200|400|500,
  "message": "description",
  "data": {
    "contact_info": {"name", "email", "phone", "links[]"},
    "education": [{"degree", "institution", "year"}],
    "experience": [{"company", "role", "start_date", "end_date", "type", "duration"}],
    "skills": {"matched": [], "missing": []},
    "projects": [{"title", "description", "tech_stack"}],
    "certifications": [{"name", "org", "year"}],
    "role_inference": "best guess candidate role",
    ... and many more structured fields
  }
}
"""

import json
import sys
import os
from typing import Dict, Any
from document_parser import parse_document
from feedback_report import FeedbackReportGenerator


def analyze_resume_api(file_path: str, 
                      job_description: str = None,
                      job_title: str = None, 
                      company: str = None,
                      required_skills: list = None,
                      preferred_skills: list = None) -> Dict[str, Any]:
    """
    Complete resume analysis API endpoint
    
    Returns structured JSON response with comprehensive analysis
    """
    
    # Validate input
    if not file_path or not os.path.exists(file_path):
        return {
            "status": "error",
            "status_code": 400,
            "message": "File not found or invalid file path",
            "data": {}
        }
    
    try:
        # Step 1: Complete document analysis
        parse_result = parse_document(
            file_path=file_path,
            job_description=job_description,
            job_title=job_title,
            company=company,
            required_skills=required_skills,
            preferred_skills=preferred_skills
        )
        
        if not parse_result.get("success"):
            return {
                "status": "error",
                "status_code": 400,
                "message": "Document parsing failed",
                "data": {
                    "error": parse_result.get("error", "Unknown parsing error"),
                    "details": parse_result
                }
            }
        
        # Step 2: Generate structured feedback report
        generator = FeedbackReportGenerator()
        feedback_report = generator.generate_feedback_report(
            parse_result=parse_result,
            job_description=job_description,
            required_skills=required_skills
        )
        
        # Step 3: Return structured API response
        return {
            "status": "success",
            "status_code": 200,
            "message": "Resume analysis completed successfully",
            "data": feedback_report
        }
        
    except Exception as e:
        return {
            "status": "error",
            "status_code": 500,
            "message": "Internal server error during analysis",
            "data": {
                "error": str(e),
                "details": {}
            }
        }


def demo_api():
    """Demo the API with a sample resume"""
    
    print("🚀 Resume Analyzer API Demo")
    print("=" * 60)
    
    # Create sample resume
    from docx import Document
    
    doc = Document()
    doc.add_heading('SARAH JOHNSON', 0)
    doc.add_paragraph('Senior Software Engineer')
    doc.add_paragraph('sarah.johnson@email.com | (555) 987-6543')
    doc.add_paragraph('LinkedIn: linkedin.com/in/sarahjohnson')
    doc.add_paragraph('')
    
    doc.add_heading('EXPERIENCE', level=1)
    doc.add_paragraph('Senior Software Engineer - TechCorp (2021-2024)')
    doc.add_paragraph('• Led development team of 8 engineers')
    doc.add_paragraph('• Built React applications serving 100K+ users daily')
    doc.add_paragraph('• Reduced deployment time by 60% using Docker/Kubernetes')
    doc.add_paragraph('• Increased code coverage from 65% to 92%')
    doc.add_paragraph('')
    
    doc.add_paragraph('Software Engineer - StartupXYZ (2019-2021)')
    doc.add_paragraph('• Developed Python microservices handling 50K+ requests/hour')
    doc.add_paragraph('• Implemented CI/CD pipelines reducing bugs by 40%')
    doc.add_paragraph('')
    
    doc.add_heading('EDUCATION', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    doc.add_paragraph('University of California, Berkeley (2015-2019)')
    doc.add_paragraph('')
    
    doc.add_heading('SKILLS', level=1)
    doc.add_paragraph('Languages: Python, JavaScript, TypeScript, Java')
    doc.add_paragraph('Frontend: React, Vue.js, HTML5, CSS3')
    doc.add_paragraph('Backend: Node.js, Django, PostgreSQL, MongoDB')
    doc.add_paragraph('DevOps: Docker, Kubernetes, AWS, CI/CD')
    
    demo_file = "demo_resume.docx"
    doc.save(demo_file)
    
    try:
        # Test API with job context
        print("📋 Analyzing resume with job context...")
        
        api_response = analyze_resume_api(
            file_path=demo_file,
            job_description="""
            We are seeking a Senior Software Engineer to lead our development team.
            
            Required Skills:
            - 5+ years of software development experience
            - Strong Python and JavaScript expertise
            - React and frontend development experience
            - Database design and optimization
            - Leadership and team management experience
            - Cloud platform experience (AWS preferred)
            
            Preferred:
            - Docker/Kubernetes containerization
            - CI/CD pipeline experience
            - Full-stack development background
            """,
            job_title="Senior Software Engineer",
            company="TechInnovate Corp",
            required_skills=["Python", "JavaScript", "React", "PostgreSQL", "Leadership", "AWS"],
            preferred_skills=["Docker", "Kubernetes", "CI/CD", "Vue.js"]
        )
        
        # Display API response
        print(f"\n✅ API Response: {api_response['status']} ({api_response['status_code']})")
        print(f"📝 Message: {api_response['message']}")
        
        if api_response['status'] == 'success':
            data = api_response['data']
            
            print(f"\n📊 **Structured JSON Response Overview:**")
            print(f"   Timestamp: {data['timestamp']}")
            print(f"   Status: {data['status']}")
            
            # Contact Information
            contact = data['contact_info']
            print(f"\n👤 **Contact Info:**")
            print(f"   Name: {contact['name']}")
            print(f"   Email: {contact['email']} (Valid: {contact['email_valid']})")
            print(f"   Phone: {contact['phone']} (Provided: {contact['phone_provided']})")
            print(f"   Professional Links: {len(contact['links'])}")
            print(f"   Completeness Score: {contact['completeness_score']}/100")
            
            # Education
            education = data['education']
            print(f"\n🎓 **Education ({len(education)} entries):**")
            for edu in education:
                print(f"   • {edu['degree']} - {edu['institution']} ({edu['year']})")
            
            # Experience
            experience = data['experience']
            print(f"\n💼 **Experience ({len(experience)} positions):**")
            for exp in experience:
                print(f"   • {exp['role']} at {exp['company']}")
                print(f"     Duration: {exp['duration']} | Type: {exp['type']}")
            
            # Skills Analysis
            skills = data['skills']
            print(f"\n🔧 **Skills Analysis:**")
            print(f"   Total Skills Found: {skills['total_count']}")
            print(f"   Required Skills Matched: {len(skills['matched'])}/{len(skills['matched']) + len(skills['missing'])}")
            print(f"   Skills Match Percentage: {skills['match_percentage']:.1f}%")
            print(f"   Missing Critical Skills: {', '.join(skills['missing']) if skills['missing'] else 'None'}")
            print(f"   Bonus Skills: {', '.join(skills['bonus'][:5])}")
            
            # Projects
            projects = data['projects']
            print(f"\n📁 **Projects ({len(projects)} projects):**")
            for project in projects:
                print(f"   • {project['title']}")
                print(f"     Tech Stack: {', '.join(project['tech_stack'][:5])}")
            
            # Certifications
            certifications = data['certifications']
            print(f"\n🏆 **Certifications ({len(certifications)} certifications):**")
            for cert in certifications:
                print(f"   • {cert['name']} - {cert['organization']} ({cert['year']})")
            
            # Role Inference
            print(f"\n🎯 **Role Inference:** {data['role_inference']}")
            
            # Quality Scores
            quality = data['quality_scores']
            print(f"\n📈 **Quality Analysis:**")
            if quality['available']:
                print(f"   Overall Score: {quality['overall_score']}/100 ({quality['quality_level']})")
                percentiles = quality['percentiles']
                print(f"   Content Fit: {percentiles['content_fit']:.1f}%")
                print(f"   Clarity & Quantification: {percentiles['clarity']:.1f}%")
                print(f"   Structure & Readability: {percentiles['structure']:.1f}%")
                print(f"   ATS Friendliness: {percentiles['ats_friendliness']:.1f}%")
            else:
                print(f"   Quality analysis not available: {quality.get('error', 'Unknown')}")
            
            # ATS Compatibility
            ats = data['ats_compatibility']
            print(f"\n🤖 **ATS Compatibility:**")
            if ats['available']:
                print(f"   ATS Score: {ats['score']}/100 ({ats['compatibility_level']})")
                print(f"   Priority Issues: {len(ats['priority_issues'])}")
                if ats['priority_issues']:
                    for issue in ats['priority_issues'][:3]:
                        print(f"     • {issue}")
            
            # Recommendations
            recommendations = data['recommendations']
            print(f"\n💡 **Top Recommendations:**")
            for i, rec in enumerate(recommendations['top_3'], 1):
                print(f"   {i}. {rec}")
            
            # Experience Summary
            exp_summary = data['experience_summary']
            print(f"\n📊 **Experience Summary:**")
            print(f"   Total Years: {exp_summary['total_years']}")
            print(f"   Career Level: {exp_summary['career_level']}")
            print(f"   Most Recent Role: {exp_summary['most_recent_role']}")
            
            # Match Analysis
            match = data['match_analysis']
            print(f"\n🎯 **Job Match Analysis:**")
            if match['available']:
                print(f"   Overall Match Score: {match['overall_match']}/100")
                print(f"   Experience Level Match: {match['experience_level_match']}")
                print(f"   Skills Coverage: {match['skills_coverage']:.1f}%")
            
            # Document Metrics
            metrics = data['document_metrics']
            content = metrics['content_metrics']
            print(f"\n📄 **Document Metrics:**")
            print(f"   File: {metrics['file_info']['filename']} ({metrics['file_info']['format']})")
            print(f"   Content: {content['total_words']} words, {content['estimated_pages']} pages")
            print(f"   Structure: {content['sections_count']} sections, {content['bullet_points']} bullets")
            
            # Save complete response
            with open("complete_api_demo.json", "w") as f:
                json.dump(api_response, f, indent=2, ensure_ascii=False)
            
            print(f"\n💾 **Complete API Response saved to:** complete_api_demo.json")
            print(f"\n🎯 **API Response Structure:**")
            print(f"   Status Code: 200 OK")
            print(f"   Top-level Keys: {list(api_response.keys())}")
            print(f"   Data Keys: {list(data.keys())}")
            
            print(f"\n✅ **API Demo completed successfully!**")
            
        else:
            print(f"❌ Analysis failed: {api_response.get('data', {}).get('error', 'Unknown error')}")
    
    except Exception as e:
        print(f"❌ Demo failed: {str(e)}")
        import traceback
        traceback.print_exc()
    
    finally:
        if os.path.exists(demo_file):
            os.remove(demo_file)
            print(f"🗑️ Cleaned up demo file")


def main():
    """Main CLI interface"""
    
    if len(sys.argv) < 2:
        print("Resume Analyzer API - Usage:")
        print("\n1. Run Demo:")
        print("   python final_api.py demo")
        print("\n2. Analyze Resume:")
        print("   python final_api.py <resume_file> [job_description] [job_title] [company] [required_skills] [preferred_skills]")
        print("\n3. Example with JSON skills:")
        print('   python final_api.py resume.pdf "job description" "Software Engineer" "TechCorp" \'["Python","React"]\' \'["Docker"]\'')
        sys.exit(1)
    
    if sys.argv[1].lower() == "demo":
        demo_api()
        return
    
    # Parse command line arguments
    file_path = sys.argv[1]
    job_description = sys.argv[2] if len(sys.argv) > 2 and sys.argv[2] != 'None' else None
    job_title = sys.argv[3] if len(sys.argv) > 3 and sys.argv[3] != 'None' else None
    company = sys.argv[4] if len(sys.argv) > 4 and sys.argv[4] != 'None' else None
    
    required_skills = None
    preferred_skills = None
    
    if len(sys.argv) > 5 and sys.argv[5] != 'None':
        try:
            required_skills = json.loads(sys.argv[5])
        except json.JSONDecodeError:
            print("Warning: Could not parse required_skills JSON")
    
    if len(sys.argv) > 6 and sys.argv[6] != 'None':
        try:
            preferred_skills = json.loads(sys.argv[6])
        except json.JSONDecodeError:
            print("Warning: Could not parse preferred_skills JSON")
    
    # Call API
    result = analyze_resume_api(
        file_path=file_path,
        job_description=job_description,
        job_title=job_title,
        company=company,
        required_skills=required_skills,
        preferred_skills=preferred_skills
    )
    
    # Output JSON response
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
